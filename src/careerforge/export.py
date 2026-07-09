"""Convert the deep agent's markdown output into plain-text and Word (.docx)
downloads.

The generated documents (resumes, cover letters, interview prep) follow
fairly simple, predictable markdown — headers, bold, bullet lists,
horizontal rules — produced by the resume-tailoring/interview-prep skills'
own formatting rules. A lightweight line-based converter is enough here; a
full markdown/HTML pipeline would be overkill for this shape of input.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_HR_RE = re.compile(r"^\s*-{3,}\s*$")
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")


def _strip_inline_markdown(text: str) -> str:
    text = _BOLD_RE.sub(r"\1", text)
    text = _ITALIC_RE.sub(r"\1", text)
    return text


def markdown_to_plain_text(md: str) -> str:
    """Strip markdown syntax down to clean, readable plain text."""
    lines: list[str] = []
    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if _HR_RE.match(line):
            continue  # horizontal rules add no value in plain text
        header_match = _HEADER_RE.match(line)
        if header_match:
            level, text = header_match.groups()
            text = _strip_inline_markdown(text).strip()
            lines.append(text.upper() if len(level) <= 2 else text)
            continue
        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            indent, content = bullet_match.groups()
            lines.append(f"{indent}- {_strip_inline_markdown(content).strip()}")
            continue
        lines.append(_strip_inline_markdown(line))
    text = re.sub(r"\n{3,}", "\n\n", "\n".join(lines))
    return text.strip() + "\n"


def _add_paragraph_with_bold(doc: Document, line: str) -> None:
    """Add a paragraph, preserving **bold** spans as actual bold runs."""
    paragraph = doc.add_paragraph()
    pos = 0
    for match in _BOLD_RE.finditer(line):
        if match.start() > pos:
            paragraph.add_run(line[pos:match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(line):
        paragraph.add_run(line[pos:])


def markdown_to_docx_bytes(md: str) -> bytes:
    """Render simple markdown into a Word document, returned as raw bytes
    suitable for ``st.download_button``."""
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip() or _HR_RE.match(line):
            continue
        header_match = _HEADER_RE.match(line)
        if header_match:
            level = min(len(header_match.group(1)), 4)
            doc.add_heading(_strip_inline_markdown(header_match.group(2)).strip(), level=level)
            continue
        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            _, content = bullet_match.groups()
            doc.add_paragraph(_strip_inline_markdown(content).strip(), style="List Bullet")
            continue
        _add_paragraph_with_bold(doc, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
