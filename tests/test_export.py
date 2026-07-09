import io

from docx import Document

from careerforge.export import markdown_to_docx_bytes, markdown_to_plain_text


def test_markdown_to_plain_text_strips_bold_and_headers():
    md = "# Niaz Mahmud\n\n**AI Engineer** | Python Developer\n\n- Built things\n- Shipped things"
    text = markdown_to_plain_text(md)
    assert "**" not in text
    assert "#" not in text
    assert "NIAZ MAHMUD" in text  # h1 uppercased
    assert "AI Engineer | Python Developer" in text
    assert "- Built things" in text


def test_markdown_to_plain_text_drops_horizontal_rules():
    md = "Summary\n\n---\n\nExperience"
    text = markdown_to_plain_text(md)
    assert "---" not in text


def test_markdown_to_plain_text_subheader_not_uppercased():
    md = "### Key Skills\nPython, SQL"
    text = markdown_to_plain_text(md)
    assert "Key Skills" in text
    assert "KEY SKILLS" not in text


def test_markdown_to_docx_bytes_produces_valid_docx():
    md = "# Niaz Mahmud\n\n**AI Engineer** | Python Developer\n\n- Built things\n- Shipped things"
    data = markdown_to_docx_bytes(md)
    assert data[:2] == b"PK"  # docx is a zip archive

    doc = Document(io.BytesIO(data))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Niaz Mahmud" in all_text
    assert "AI Engineer" in all_text
    assert "Built things" in all_text


def test_markdown_to_docx_preserves_bold_run():
    data = markdown_to_docx_bytes("**Bold** and normal")
    doc = Document(io.BytesIO(data))
    runs = doc.paragraphs[0].runs
    bold_runs = [r for r in runs if r.bold]
    assert any(r.text == "Bold" for r in bold_runs)
